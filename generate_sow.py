#!/usr/bin/env python3
"""
Generate Statement of Work (SoW) document for Sparking - Smart Parking Management System
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Circle, Rectangle
import numpy as np

# Output directory
OUTPUT_DIR = "/Users/sudipto/Desktop/projects/sparking/sow_documents"
os.makedirs(OUTPUT_DIR, exist_ok=True)

def create_data_flow_diagram():
    """Create a comprehensive data flow diagram"""
    fig, ax = plt.subplots(1, 1, figsize=(16, 12))
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 12)
    ax.set_aspect('equal')
    ax.axis('off')

    # Title
    ax.text(8, 11.5, 'SPARKING - Data Flow Diagram', fontsize=18, fontweight='bold',
            ha='center', va='center', color='#1a365d')

    # Color scheme
    colors = {
        'external': '#E3F2FD',  # Light blue - external entities
        'process': '#E8F5E9',   # Light green - processes
        'database': '#FFF3E0',  # Light orange - data stores
        'ai': '#F3E5F5',        # Light purple - AI components
        'border': '#37474F'
    }

    def draw_box(x, y, w, h, label, color, fontsize=9, bold=False):
        rect = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.1",
                              facecolor=color, edgecolor=colors['border'], linewidth=1.5)
        ax.add_patch(rect)
        weight = 'bold' if bold else 'normal'
        ax.text(x + w/2, y + h/2, label, ha='center', va='center', fontsize=fontsize,
                fontweight=weight, wrap=True)

    def draw_cylinder(x, y, w, h, label):
        # Database cylinder
        ellipse_h = h * 0.15
        rect = Rectangle((x, y + ellipse_h/2), w, h - ellipse_h,
                         facecolor=colors['database'], edgecolor=colors['border'], linewidth=1.5)
        ax.add_patch(rect)
        # Top ellipse
        ellipse_top = mpatches.Ellipse((x + w/2, y + h - ellipse_h/2), w, ellipse_h,
                                       facecolor=colors['database'], edgecolor=colors['border'], linewidth=1.5)
        ax.add_patch(ellipse_top)
        # Bottom ellipse
        ellipse_bottom = mpatches.Ellipse((x + w/2, y + ellipse_h/2), w, ellipse_h,
                                          facecolor=colors['database'], edgecolor=colors['border'], linewidth=1.5)
        ax.add_patch(ellipse_bottom)
        ax.text(x + w/2, y + h/2, label, ha='center', va='center', fontsize=8, fontweight='bold')

    def draw_arrow(start, end, label='', color='#455A64', curved=False):
        style = "Simple, tail_width=0.3, head_width=4, head_length=4"
        if curved:
            arrow = FancyArrowPatch(start, end, connectionstyle="arc3,rad=0.2",
                                   arrowstyle=style, color=color, linewidth=1)
        else:
            arrow = FancyArrowPatch(start, end, arrowstyle=style, color=color, linewidth=1)
        ax.add_patch(arrow)
        if label:
            mid_x = (start[0] + end[0]) / 2
            mid_y = (start[1] + end[1]) / 2
            ax.text(mid_x, mid_y + 0.2, label, fontsize=7, ha='center', color='#37474F')

    # External Entities (Row 1)
    draw_box(0.5, 9.5, 2, 1.2, 'Vehicle\nOwner/Driver', colors['external'], bold=True)
    draw_box(3.5, 9.5, 2, 1.2, 'Kiosk\nTerminal', colors['external'], bold=True)
    draw_box(6.5, 9.5, 2, 1.2, 'Mobile App\n(Future)', colors['external'], bold=True)
    draw_box(9.5, 9.5, 2, 1.2, 'Parking\nOperator', colors['external'], bold=True)
    draw_box(12.5, 9.5, 2.5, 1.2, 'System\nAdministrator', colors['external'], bold=True)

    # Core Processes (Row 2)
    draw_box(0.5, 7, 2.5, 1.3, '1.0\nEntry/Exit\nManagement', colors['process'], bold=True)
    draw_box(4, 7, 2.5, 1.3, '2.0\nToken\nGeneration', colors['process'], bold=True)
    draw_box(7.5, 7, 2.5, 1.3, '3.0\nSlot\nAllocation', colors['process'], bold=True)
    draw_box(11, 7, 2.5, 1.3, '4.0\nPayment\nProcessing', colors['process'], bold=True)

    # AI & Real-time Layer (Row 3)
    draw_box(0.5, 4.5, 3, 1.3, '5.0\nAI Vehicle/Plate\nDetection', colors['ai'], bold=True)
    draw_box(4.5, 4.5, 3, 1.3, '6.0\nReal-time\nOccupancy Tracking', colors['ai'], bold=True)
    draw_box(8.5, 4.5, 3, 1.3, '7.0\nWallet & Transaction\nManagement', colors['process'], bold=True)
    draw_box(12.5, 4.5, 2.5, 1.3, '8.0\nAnalytics &\nReporting', colors['process'], bold=True)

    # Hardware Layer (Row 4)
    draw_box(0.5, 2.5, 2, 1, 'CCTV\nCameras', colors['external'], fontsize=8)
    draw_box(3, 2.5, 2, 1, 'Entry/Exit\nGates', colors['external'], fontsize=8)
    draw_box(5.5, 2.5, 2, 1, 'Display\nBoards', colors['external'], fontsize=8)
    draw_box(8, 2.5, 2, 1, 'Payment\nGateways', colors['external'], fontsize=8)
    draw_box(10.5, 2.5, 2, 1, 'Bank\nAPIs', colors['external'], fontsize=8)
    draw_box(13, 2.5, 2, 1, 'Email/SMS\nServices', colors['external'], fontsize=8)

    # Data Stores (Bottom)
    draw_cylinder(1, 0.3, 2.5, 1.5, 'PostgreSQL\nDatabase')
    draw_cylinder(4.5, 0.3, 2, 1.5, 'Redis\nCache')
    draw_cylinder(7.5, 0.3, 2.5, 1.5, 'File\nStorage')
    draw_cylinder(11, 0.3, 2.5, 1.5, 'Audit\nLogs')

    # Draw arrows - External to Processes
    draw_arrow((1.5, 9.5), (1.5, 8.3), 'Entry Request')
    draw_arrow((4.5, 9.5), (5, 8.3), 'Token Request')
    draw_arrow((10.5, 9.5), (10, 8.3), 'Dashboard')
    draw_arrow((13.75, 9.5), (12.5, 8.3), 'Config')

    # Processes to each other
    draw_arrow((3, 7.65), (4, 7.65), '')
    draw_arrow((6.5, 7.65), (7.5, 7.65), '')
    draw_arrow((10, 7.65), (11, 7.65), '')

    # Processes to AI Layer
    draw_arrow((1.75, 7), (2, 5.8), '')
    draw_arrow((5.25, 7), (6, 5.8), '')
    draw_arrow((8.75, 7), (8.75, 5.8), '')
    draw_arrow((12.25, 7), (13.75, 5.8), '')

    # AI to Hardware
    draw_arrow((2, 4.5), (1.5, 3.5), 'Video Feed')
    draw_arrow((3.5, 5.1), (4, 3.5), 'Gate Control')
    draw_arrow((6, 4.5), (6.5, 3.5), 'Display Data')
    draw_arrow((10, 4.5), (9, 3.5), 'Payment')
    draw_arrow((10, 5.1), (11.5, 3.5), 'Bank Transfer')
    draw_arrow((14, 4.5), (14, 3.5), 'Alerts')

    # To Data Stores
    draw_arrow((2.25, 2.5), (2.25, 1.8), '')
    draw_arrow((5.5, 2.5), (5.5, 1.8), '')
    draw_arrow((8.75, 2.5), (8.75, 1.8), '')
    draw_arrow((12.25, 4.5), (12.25, 1.8), '')

    # Legend
    legend_y = 0.3
    draw_box(14, 2, 1.5, 0.4, 'External Entity', colors['external'], fontsize=7)
    draw_box(14, 1.5, 1.5, 0.4, 'Process', colors['process'], fontsize=7)
    draw_box(14, 1, 1.5, 0.4, 'AI Component', colors['ai'], fontsize=7)
    draw_box(14, 0.5, 1.5, 0.4, 'Data Store', colors['database'], fontsize=7)

    plt.tight_layout()
    filepath = os.path.join(OUTPUT_DIR, 'data_flow_diagram.png')
    plt.savefig(filepath, dpi=150, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close()
    print(f"Created: {filepath}")
    return filepath


def create_architecture_diagram():
    """Create a software architecture diagram"""
    fig, ax = plt.subplots(1, 1, figsize=(18, 16))
    ax.set_xlim(0, 18)
    ax.set_ylim(0, 16)
    ax.set_aspect('equal')
    ax.axis('off')

    # Title
    ax.text(9, 15.3, 'SPARKING - Software Architecture Diagram', fontsize=20, fontweight='bold',
            ha='center', va='center', color='#1a365d')

    # Color scheme for layers
    colors = {
        'client': '#BBDEFB',      # Blue - Client Layer
        'presentation': '#C8E6C9', # Green - Presentation
        'api': '#FFE0B2',         # Orange - API Layer
        'business': '#E1BEE7',    # Purple - Business Logic
        'data': '#FFCCBC',        # Red - Data Layer
        'external': '#F5F5F5',    # Gray - External
        'infra': '#B2DFDB',       # Teal - Infrastructure
        'ai': '#F3E5F5',          # Light purple - AI
        'border': '#37474F'
    }

    def draw_layer(y, height, color, label, width=14):
        rect = FancyBboxPatch((0.5, y), width, height, boxstyle="round,pad=0.02,rounding_size=0.2",
                              facecolor=color, edgecolor=colors['border'], linewidth=2, alpha=0.7)
        ax.add_patch(rect)
        ax.text(0.9, y + height - 0.35, label, fontsize=10, fontweight='bold',
                va='top', color='#1a365d')

    def draw_component(x, y, w, h, label, sublabel='', color='white', fontsize_main=8, fontsize_sub=6):
        rect = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.1",
                              facecolor=color, edgecolor=colors['border'], linewidth=1.5)
        ax.add_patch(rect)
        if sublabel:
            ax.text(x + w/2, y + h/2 + 0.2, label, ha='center', va='center',
                    fontsize=fontsize_main, fontweight='bold')
            ax.text(x + w/2, y + h/2 - 0.25, sublabel, ha='center', va='center',
                    fontsize=fontsize_sub, color='#616161')
        else:
            ax.text(x + w/2, y + h/2, label, ha='center', va='center',
                    fontsize=fontsize_main, fontweight='bold')

    # Layer 1: Client Layer (Top)
    draw_layer(13, 1.8, colors['client'], 'CLIENT LAYER')
    draw_component(1.5, 13.35, 2.8, 1.1, 'Web Dashboard', 'React/Next.js')
    draw_component(4.8, 13.35, 2.8, 1.1, 'Public Kiosk', 'Touch Interface')
    draw_component(8.1, 13.35, 2.8, 1.1, 'Find My Car', 'Public Portal')
    draw_component(11.4, 13.35, 2.8, 1.1, 'Admin Panel', 'Settings/Config')

    # Layer 2: Presentation Layer
    draw_layer(10.8, 1.9, colors['presentation'], 'PRESENTATION LAYER (Next.js App Router)')
    draw_component(1.2, 11.15, 2.4, 1.2, 'Server\nComponents', '')
    draw_component(4, 11.15, 2.4, 1.2, 'Client\nComponents', '')
    draw_component(6.8, 11.15, 2.4, 1.2, 'Layouts &\nPages', '')
    draw_component(9.6, 11.15, 2.4, 1.2, 'React Query\nState', '')
    draw_component(12.4, 11.15, 2, 1.2, 'Zustand\nStore', '')

    # Layer 3: API Layer
    draw_layer(8.4, 2.1, colors['api'], 'API LAYER (RESTful + GraphQL)')
    draw_component(1.2, 8.8, 2.5, 1.4, 'Auth API', '/api/auth/*')
    draw_component(4, 8.8, 2.5, 1.4, 'Parking API', '/api/lots/*')
    draw_component(6.8, 8.8, 2.5, 1.4, 'Token API', '/api/tokens/*')
    draw_component(9.6, 8.8, 2.5, 1.4, 'Payment API', '/api/pay/*')
    draw_component(12.4, 8.8, 2, 1.4, 'Analytics', '/api/stats/*')

    # Layer 4: Business Logic Layer
    draw_layer(5.8, 2.3, colors['business'], 'BUSINESS LOGIC LAYER')
    draw_component(1, 6.25, 2.5, 1.5, 'Slot Allocation\nEngine', 'Smart Assign')
    draw_component(3.8, 6.25, 2.5, 1.5, 'Pricing\nEngine', 'Dynamic Price')
    draw_component(6.6, 6.25, 2.5, 1.5, 'Wallet\nService', 'Transactions')
    draw_component(9.4, 6.25, 2.5, 1.5, 'Notification\nService', 'Email/SMS')
    draw_component(12.2, 6.25, 2.2, 1.5, 'Analytics\nEngine', 'Reports')

    # Layer 5: Data Access Layer
    draw_layer(3.3, 2.2, colors['data'], 'DATA ACCESS LAYER')
    draw_component(1.5, 3.7, 2.8, 1.5, 'Prisma ORM', 'PostgreSQL')
    draw_component(4.8, 3.7, 2.8, 1.5, 'Cache Layer', 'Redis Client')
    draw_component(8.1, 3.7, 2.8, 1.5, 'File Storage', 'Local / S3')
    draw_component(11.4, 3.7, 2.8, 1.5, 'Encryption', 'Credentials')

    # Layer 6: Infrastructure Layer (Bottom)
    draw_layer(0.8, 2.2, colors['infra'], 'INFRASTRUCTURE LAYER')
    draw_component(1, 1.2, 2, 1.4, 'PostgreSQL', 'Primary DB')
    draw_component(3.3, 1.2, 2, 1.4, 'Redis', 'Cache')
    draw_component(5.6, 1.2, 2, 1.4, 'MQTT', 'Mosquitto')
    draw_component(7.9, 1.2, 2, 1.4, 'Socket.IO', 'Real-time')
    draw_component(10.2, 1.2, 2, 1.4, 'Docker', 'Containers')
    draw_component(12.5, 1.2, 2, 1.4, 'Nginx', 'Proxy')

    # AI Pipeline (Side box) - positioned to the right
    ai_box = FancyBboxPatch((15, 5.8), 2.5, 5.5, boxstyle="round,pad=0.02,rounding_size=0.15",
                            facecolor=colors['ai'], edgecolor=colors['border'], linewidth=2)
    ax.add_patch(ai_box)
    ax.text(16.25, 10.8, 'AI PIPELINE', fontsize=9, fontweight='bold', ha='center', color='#1a365d')
    ax.text(16.25, 9.8, 'Intel', fontsize=8, ha='center', fontweight='bold')
    ax.text(16.25, 9.3, 'OpenVINO', fontsize=8, ha='center')
    ax.text(16.25, 8.3, 'YOLOv8', fontsize=8, ha='center', fontweight='bold')
    ax.text(16.25, 7.8, 'Detection', fontsize=8, ha='center')
    ax.text(16.25, 6.8, 'License Plate', fontsize=8, ha='center', fontweight='bold')
    ax.text(16.25, 6.3, 'Recognition', fontsize=8, ha='center')

    # External Services box - positioned to the right bottom
    ext_box = FancyBboxPatch((15, 0.8), 2.5, 4.5, boxstyle="round,pad=0.02,rounding_size=0.15",
                             facecolor='#ECEFF1', edgecolor=colors['border'], linewidth=2)
    ax.add_patch(ext_box)
    ax.text(16.25, 4.8, 'EXTERNAL', fontsize=9, fontweight='bold', ha='center', color='#1a365d')
    ax.text(16.25, 4.3, 'SERVICES', fontsize=9, fontweight='bold', ha='center', color='#1a365d')
    ax.text(16.25, 3.5, 'Razorpay', fontsize=8, ha='center')
    ax.text(16.25, 3.0, 'Stripe', fontsize=8, ha='center')
    ax.text(16.25, 2.5, 'Twilio/MSG91', fontsize=7, ha='center')
    ax.text(16.25, 2.0, 'SendGrid', fontsize=8, ha='center')
    ax.text(16.25, 1.5, 'AWS S3', fontsize=8, ha='center')

    # Arrows between layers (vertical flow)
    arrow_x_positions = [2.9, 6.2, 9.5, 12.8]
    for x in arrow_x_positions:
        # Client to Presentation
        ax.annotate('', xy=(x, 13), xytext=(x, 12.7),
                   arrowprops=dict(arrowstyle='->', color='#455A64', lw=1.2))
        # Presentation to API
        ax.annotate('', xy=(x, 10.8), xytext=(x, 10.5),
                   arrowprops=dict(arrowstyle='->', color='#455A64', lw=1.2))
        # API to Business
        ax.annotate('', xy=(x, 8.4), xytext=(x, 8.1),
                   arrowprops=dict(arrowstyle='->', color='#455A64', lw=1.2))
        # Business to Data
        ax.annotate('', xy=(x, 5.8), xytext=(x, 5.5),
                   arrowprops=dict(arrowstyle='->', color='#455A64', lw=1.2))
        # Data to Infrastructure
        ax.annotate('', xy=(x, 3.3), xytext=(x, 3.0),
                   arrowprops=dict(arrowstyle='->', color='#455A64', lw=1.2))

    # Arrow from API layer to AI Pipeline
    ax.annotate('', xy=(15, 8.5), xytext=(14.5, 8.5),
               arrowprops=dict(arrowstyle='<->', color='#7B1FA2', lw=1.5))

    # Arrow from Data layer to External Services
    ax.annotate('', xy=(15, 3.5), xytext=(14.5, 3.5),
               arrowprops=dict(arrowstyle='<->', color='#455A64', lw=1.5))

    plt.tight_layout()
    filepath = os.path.join(OUTPUT_DIR, 'software_architecture_diagram.png')
    plt.savefig(filepath, dpi=150, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close()
    print(f"Created: {filepath}")
    return filepath


def set_cell_shading(cell, color):
    """Set background color for a table cell"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def create_sow_document(data_flow_path, architecture_path):
    """Create the Statement of Work document"""
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ========== COVER PAGE ==========
    # Add spacing before title
    for _ in range(6):
        doc.add_paragraph()

    # Title
    title = doc.add_paragraph()
    title_run = title.add_run('STATEMENT OF WORK')
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(26, 54, 93)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('SPARKING')
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(36)
    subtitle_run.font.color.rgb = RGBColor(46, 125, 50)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Tagline
    tagline = doc.add_paragraph()
    tagline_run = tagline.add_run('AI-Powered Smart Parking Management System')
    tagline_run.font.size = Pt(16)
    tagline_run.font.color.rgb = RGBColor(97, 97, 97)
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add spacing
    for _ in range(4):
        doc.add_paragraph()

    # Version info
    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_run = version_para.add_run('Version 1.0\nJanuary 2026')
    version_run.font.size = Pt(12)
    version_run.font.color.rgb = RGBColor(117, 117, 117)

    # Page break
    doc.add_page_break()

    # ========== TABLE OF CONTENTS ==========
    toc_heading = doc.add_heading('Table of Contents', level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add some spacing after heading
    doc.add_paragraph()

    toc_items = [
        ('1.', 'Executive Summary', '3'),
        ('2.', 'Feature List', '4'),
        ('3.', 'Technology Stack', '8'),
        ('4.', 'Data Flow Diagram', '11'),
        ('5.', 'Software Architecture', '12'),
        ('6.', 'Plan of Action (Phase-wise Development)', '13'),
        ('7.', 'Required Team Setup', '16'),
        ('8.', 'Project Timeline', '18'),
        ('9.', 'API Documentation', '20'),
        ('10.', 'Database Schema', '25'),
    ]

    # Create a properly formatted TOC table with column widths
    toc_table = doc.add_table(rows=len(toc_items), cols=3)
    toc_table.autofit = False

    # Set column widths (number, title, page)
    for row in toc_table.rows:
        row.cells[0].width = Cm(1.5)
        row.cells[1].width = Cm(12)
        row.cells[2].width = Cm(2)

    for i, (num, title, page) in enumerate(toc_items):
        row = toc_table.rows[i]

        # Section number - left aligned, bold
        cell0 = row.cells[0]
        cell0.text = num
        cell0.paragraphs[0].runs[0].font.bold = True
        cell0.paragraphs[0].runs[0].font.size = Pt(12)
        cell0.vertical_alignment = 1  # CENTER

        # Title - left aligned
        cell1 = row.cells[1]
        cell1.text = title
        cell1.paragraphs[0].runs[0].font.size = Pt(12)
        cell1.vertical_alignment = 1  # CENTER

        # Page number - right aligned
        cell2 = row.cells[2]
        cell2.text = page
        cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cell2.paragraphs[0].runs[0].font.size = Pt(12)
        cell2.vertical_alignment = 1  # CENTER

        # Add spacing between rows
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(6)
            cell.paragraphs[0].paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ========== 1. EXECUTIVE SUMMARY ==========
    doc.add_heading('1. Executive Summary', level=1)

    exec_summary = """
SPARKING is a comprehensive, enterprise-grade Smart Parking Management System designed to revolutionize parking facility operations through artificial intelligence and real-time monitoring. The system combines cutting-edge computer vision technology with a robust payment infrastructure to deliver a seamless parking experience for both operators and end-users.

The platform addresses critical challenges faced by modern parking facilities including:
• Real-time occupancy tracking and management
• Automated vehicle detection and license plate recognition
• Intelligent slot allocation and guidance
• Secure, multi-method payment processing
• Comprehensive analytics and revenue management

Target Markets:
• Airports and Transportation Hubs
• Shopping Malls and Retail Centers
• Hospitals and Medical Facilities
• Commercial Office Buildings
• Hotels and Hospitality Venues
• Sports Stadiums and Event Venues
• Residential Complexes

Key Differentiators:
1. AI-Powered Detection: Utilizes Intel OpenVINO and YOLOv8 for real-time vehicle and license plate detection with 95%+ accuracy
2. Scalable Architecture: Microservices-based design supporting multiple venues with thousands of parking slots
3. Integrated Wallet System: PayPal-like digital wallet with P2P transfers, KYC verification, and multi-gateway support
4. Real-time Operations: Socket.IO and MQTT-based real-time updates for instant occupancy changes
5. Hardware Agnostic: Compatible with existing CCTV infrastructure through RTSP/ONVIF protocols
"""

    for para in exec_summary.strip().split('\n\n'):
        p = doc.add_paragraph(para.strip())
        p.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # ========== 2. FEATURE LIST ==========
    doc.add_heading('2. Feature List', level=1)

    # 2.1 Core Parking Management
    doc.add_heading('2.1 Core Parking Management', level=2)

    core_features = [
        ('Multi-Venue Support', 'Organize and manage multiple parking facilities (airports, malls, hospitals, stadiums, etc.) from a single dashboard with independent configurations'),
        ('Multi-Level Architecture', 'Support for multi-story parking structures with floor-wise and zone-wise organization'),
        ('Zone Management', 'Create specialized zones: General, VIP, EV Charging, Disabled, Staff, Visitor, Short-term, Long-term, Two-wheeler, Valet, Reserved'),
        ('Slot Configuration', 'Individual slot setup with type (Standard, Compact, Large, Handicapped, EV, Motorcycle, VIP, Reserved), vehicle restrictions, and visual positioning'),
        ('Real-time Status', 'Live tracking of slot status: Available, Occupied, Reserved, Maintenance, Blocked'),
        ('Visual Slot Mapping', 'Interactive floor plans with slot positions, rotations, and live occupancy indicators'),
    ]

    for title, desc in core_features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title}: ').bold = True
        p.add_run(desc)

    # 2.2 AI-Powered Detection
    doc.add_heading('2.2 AI-Powered Detection', level=2)

    ai_features = [
        ('Vehicle Detection', 'Real-time detection using YOLOv8n model optimized for parking scenarios with 95%+ accuracy'),
        ('License Plate Recognition', 'Automatic Number Plate Recognition (ANPR) supporting multiple plate formats and regions'),
        ('Occupancy Detection', 'AI-based slot occupancy detection through bounding box matching with configurable confidence thresholds'),
        ('Multi-Camera Support', 'Parallel processing of multiple camera feeds with Intel OpenVINO acceleration'),
        ('Edge Processing', 'On-premise AI inference for reduced latency and enhanced privacy'),
        ('Detection Events', 'Comprehensive logging of all detection events with timestamps, confidence scores, and associated metadata'),
    ]

    for title, desc in ai_features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title}: ').bold = True
        p.add_run(desc)

    # 2.3 Entry/Exit Management
    doc.add_heading('2.3 Entry/Exit Management', level=2)

    entry_features = [
        ('Token Generation', 'Multiple token types: QR Code (auto-generated), RFID, Barcode, ANPR-based, Manual entry'),
        ('Smart Slot Allocation', 'Intelligent algorithm considering vehicle type, zone preference, accessibility needs, and EV charging requirements'),
        ('Gate Control Integration', 'Support for Entry, Exit, and Bidirectional gates with RS485, Relay, and API-based hardware'),
        ('Automatic Gate Opening', 'Token validation triggers automatic gate operation with manual override capability'),
        ('Visit History', 'Complete vehicle visit tracking with entry/exit timestamps and duration calculations'),
    ]

    for title, desc in entry_features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title}: ').bold = True
        p.add_run(desc)

    # 2.4 Payment & Wallet System
    doc.add_heading('2.4 Payment & Wallet System', level=2)

    payment_features = [
        ('Digital Wallet', 'Built-in wallet system with Personal, Business, and Merchant account types'),
        ('Multiple Pricing Models', 'Flat Rate, Hourly (with peak multipliers), Slab-based, Dynamic pricing, Free parking options'),
        ('Payment Methods', 'Cash, Card (Stripe), UPI (Razorpay), Wallet balance, Postpaid accounts'),
        ('P2P Transfers', 'Peer-to-peer wallet transfers between users with transaction limits'),
        ('Bank Integration', 'Link bank accounts for deposits/withdrawals with penny-drop verification'),
        ('KYC Verification', 'Tiered KYC levels (None, Basic, Intermediate, Full) with corresponding transaction limits'),
        ('Transaction Management', 'Complete transaction history with statuses: Pending, Processing, Completed, Failed, Cancelled, Reversed'),
        ('Refund Processing', 'Automated and manual refund capabilities with audit trail'),
    ]

    for title, desc in payment_features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title}: ').bold = True
        p.add_run(desc)

    # 2.5 Analytics & Reporting
    doc.add_heading('2.5 Analytics & Reporting', level=2)

    analytics_features = [
        ('Real-time Dashboard', 'Live metrics: current occupancy, available slots, active tokens, today\'s revenue, camera status'),
        ('Historical Analytics', 'Occupancy trends, revenue analysis, traffic patterns by hour/day/month'),
        ('Zone Performance', 'Comparative analysis across zones with utilization metrics'),
        ('Peak Hour Analysis', 'Identification of busiest periods for staffing and pricing optimization'),
        ('Vehicle Distribution', 'Statistics by vehicle type and parking duration'),
        ('Predictive Analytics', 'Occupancy forecasting based on historical patterns'),
        ('Custom Reports', 'Export to CSV and PDF with custom date ranges and filters'),
    ]

    for title, desc in analytics_features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title}: ').bold = True
        p.add_run(desc)

    # 2.6 User Management
    doc.add_heading('2.6 User Management & Access Control', level=2)

    user_features = [
        ('Role-Based Access', 'Five user roles: Super Admin, Admin, Operator, Auditor, Viewer'),
        ('Multi-Session Support', 'Configurable concurrent session limits per user'),
        ('Parking Lot Assignments', 'Operators can be assigned to specific facilities'),
        ('Audit Logging', 'Comprehensive tracking of all user actions with before/after values'),
        ('Session Security', 'IP and user agent tracking, automatic session cleanup'),
    ]

    for title, desc in user_features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title}: ').bold = True
        p.add_run(desc)

    # 2.7 Notifications & Alerts
    doc.add_heading('2.7 Notifications & Alerts', level=2)

    notification_features = [
        ('Multi-Channel Delivery', 'In-app notifications, Email (SMTP/Resend/SendGrid), SMS (Twilio/MSG91)'),
        ('Alert Rules', 'Configurable metric-based alerts with operators (GT, LT, EQ, GTE, LTE)'),
        ('Alert Actions', 'Email, SMS, Push notifications, Webhook triggers'),
        ('Cooldown Periods', 'Prevent alert fatigue with configurable cooldown intervals'),
        ('Notification Types', 'Payment confirmations, session alerts, overstay warnings, system notifications'),
    ]

    for title, desc in notification_features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title}: ').bold = True
        p.add_run(desc)

    # 2.8 Hardware Integration
    doc.add_heading('2.8 Hardware Integration', level=2)

    hardware_features = [
        ('Camera Support', 'RTSP and ONVIF protocol support with PTZ and IR capabilities'),
        ('Display Integration', 'LED counters, LCD signage, kiosks, directional displays'),
        ('Gate Systems', 'Entry/exit gates with multiple control protocols'),
        ('Health Monitoring', 'Real-time camera status tracking (Online, Offline, Error, Maintenance)'),
        ('Encrypted Credentials', 'Secure storage of camera and hardware credentials'),
    ]

    for title, desc in hardware_features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title}: ').bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ========== 3. TECHNOLOGY STACK ==========
    doc.add_heading('3. Technology Stack', level=1)

    # Frontend
    doc.add_heading('3.1 Frontend Technologies', level=2)

    frontend_table = doc.add_table(rows=1, cols=3)
    frontend_table.style = 'Table Grid'

    header_cells = frontend_table.rows[0].cells
    headers = ['Category', 'Technology', 'Version/Details']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '1a365d')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    frontend_data = [
        ('Framework', 'Next.js', '16.1.1 (App Router, Server Components)'),
        ('UI Library', 'React', '19.2.3'),
        ('Styling', 'Tailwind CSS', '4.0'),
        ('Component Library', 'Radix UI', 'Full suite (15+ components)'),
        ('Forms', 'React Hook Form + Zod', '7.69.0 / 4.2.1'),
        ('State Management', 'Zustand', '5.0.9'),
        ('Data Fetching', 'TanStack React Query', '5.90.12'),
        ('Charts', 'Recharts', '2.15.4'),
        ('Maps', 'Leaflet + React-Leaflet', '1.9.4 / 5.0.0'),
        ('Icons', 'Lucide React', '0.562.0'),
        ('PDF Generation', 'jsPDF + Autotable', '3.0.4'),
        ('Real-time', 'Socket.IO Client', '4.8.3'),
        ('Language', 'TypeScript', '5.x'),
    ]

    for category, tech, version in frontend_data:
        row = frontend_table.add_row()
        row.cells[0].text = category
        row.cells[1].text = tech
        row.cells[2].text = version

    # Backend
    doc.add_heading('3.2 Backend Technologies', level=2)

    backend_table = doc.add_table(rows=1, cols=3)
    backend_table.style = 'Table Grid'

    header_cells = backend_table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '2e7d32')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    backend_data = [
        ('Runtime', 'Node.js', 'LTS (18.x+)'),
        ('Framework', 'Next.js API Routes', '16.1.1'),
        ('ORM', 'Prisma', '7.2.0'),
        ('Database', 'PostgreSQL', '16+'),
        ('Cache', 'Redis', '7.x'),
        ('Auth', 'jose (JWT) + bcryptjs', '6.1.3 / 3.0.3'),
        ('Real-time', 'Socket.IO Server', '4.8.3'),
        ('Email', 'Nodemailer', '7.0.12'),
        ('Validation', 'Zod', '4.2.1'),
    ]

    for category, tech, version in backend_data:
        row = backend_table.add_row()
        row.cells[0].text = category
        row.cells[1].text = tech
        row.cells[2].text = version

    # AI Pipeline
    doc.add_heading('3.3 AI/ML Pipeline', level=2)

    ai_table = doc.add_table(rows=1, cols=3)
    ai_table.style = 'Table Grid'

    header_cells = ai_table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '6a1b9a')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    ai_data = [
        ('Language', 'Python', '3.9+'),
        ('Inference Engine', 'Intel OpenVINO', '2024.0.0+'),
        ('Object Detection', 'Ultralytics YOLOv8', '8.0.0+'),
        ('Computer Vision', 'OpenCV', '4.8.0+'),
        ('Video Processing', 'PyAV', '10.0.0'),
        ('Messaging', 'paho-mqtt', '1.6.1'),
        ('HTTP Client', 'aiohttp + requests', '3.9.0 / 2.31.0'),
        ('Geometry', 'Shapely', '2.0.0'),
        ('Logging', 'structlog', '23.0.0'),
    ]

    for category, tech, version in ai_data:
        row = ai_table.add_row()
        row.cells[0].text = category
        row.cells[1].text = tech
        row.cells[2].text = version

    # DevOps
    doc.add_heading('3.4 DevOps & Infrastructure', level=2)

    devops_table = doc.add_table(rows=1, cols=3)
    devops_table.style = 'Table Grid'

    header_cells = devops_table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], 'e65100')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    devops_data = [
        ('Containerization', 'Docker', 'Latest'),
        ('Orchestration', 'Docker Compose', 'Multi-service setup'),
        ('Reverse Proxy', 'Nginx', 'Alpine'),
        ('Message Broker', 'Eclipse Mosquitto', '2.x (MQTT)'),
        ('Hosting', 'Vercel / Self-hosted', 'Serverless + Docker'),
        ('CI/CD', 'Git-based deployment', 'Vercel auto-deploy'),
    ]

    for category, tech, version in devops_data:
        row = devops_table.add_row()
        row.cells[0].text = category
        row.cells[1].text = tech
        row.cells[2].text = version

    # External Services
    doc.add_heading('3.5 External Services & Integrations', level=2)

    services_table = doc.add_table(rows=1, cols=3)
    services_table.style = 'Table Grid'

    header_cells = services_table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '00695c')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    services_data = [
        ('Payment - India', 'Razorpay', 'Cards, UPI, Wallets'),
        ('Payment - Global', 'Stripe', 'Cards, Apple Pay, Google Pay'),
        ('Email', 'SMTP / Resend / SendGrid', 'Transactional emails'),
        ('SMS - India', 'MSG91', 'OTP, Notifications'),
        ('SMS - Global', 'Twilio', 'SMS worldwide'),
        ('Storage', 'Local / AWS S3', 'File storage'),
    ]

    for category, tech, version in services_data:
        row = services_table.add_row()
        row.cells[0].text = category
        row.cells[1].text = tech
        row.cells[2].text = version

    doc.add_page_break()

    # ========== 4. DATA FLOW DIAGRAM ==========
    doc.add_heading('4. Data Flow Diagram', level=1)

    p = doc.add_paragraph()
    p.add_run('The following diagram illustrates the flow of data through the SPARKING system, showing how information moves between external entities, processes, and data stores.')
    p.paragraph_format.space_after = Pt(12)

    # Add the data flow diagram
    doc.add_picture(data_flow_path, width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Diagram explanation
    doc.add_heading('4.1 Data Flow Description', level=2)

    flow_descriptions = [
        ('Entry Flow', 'Vehicle arrives → Camera detects vehicle/plate → Token generated → Slot allocated → Gate opens → Occupancy updated'),
        ('Payment Flow', 'Token validated → Duration calculated → Price computed → Payment processed → Wallet debited → Receipt generated'),
        ('Exit Flow', 'Vehicle at exit → Token validated → Payment confirmed → Gate opens → Slot released → Analytics updated'),
        ('AI Detection Flow', 'Camera feed → OpenVINO inference → Detection events → MQTT/API → Database → Real-time broadcast'),
        ('Analytics Flow', 'Raw data → Aggregation jobs → Metrics calculation → Dashboard display → Report generation'),
    ]

    for title, desc in flow_descriptions:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title}: ').bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ========== 5. SOFTWARE ARCHITECTURE ==========
    doc.add_heading('5. Software Architecture', level=1)

    p = doc.add_paragraph()
    p.add_run('SPARKING follows a layered architecture pattern with clear separation of concerns. The diagram below shows the system\'s architectural components and their relationships.')
    p.paragraph_format.space_after = Pt(12)

    # Add the architecture diagram
    doc.add_picture(architecture_path, width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Architecture explanation
    doc.add_heading('5.1 Architecture Layers', level=2)

    layers = [
        ('Client Layer', 'Web dashboard, public kiosk interface, vehicle finder portal, and admin panel built with React/Next.js'),
        ('Presentation Layer', 'Next.js App Router with Server Components, Client Components, and state management via React Query and Zustand'),
        ('API Layer', 'RESTful endpoints organized by domain (auth, parking, tokens, payments, analytics) with optional GraphQL support'),
        ('Business Logic Layer', 'Core services including slot allocation engine, pricing engine, wallet service, notification service, and analytics engine'),
        ('Data Access Layer', 'Prisma ORM for PostgreSQL, Redis client for caching, file storage abstraction, and encryption services'),
        ('Infrastructure Layer', 'PostgreSQL database, Redis cache, MQTT broker, Socket.IO server, Docker containers, and Nginx reverse proxy'),
    ]

    for title, desc in layers:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title}: ').bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ========== 6. PLAN OF ACTION ==========
    doc.add_heading('6. Plan of Action (Phase-wise Development)', level=1)

    # Phase 1
    doc.add_heading('Phase 1: Foundation & Core Infrastructure', level=2)
    p = doc.add_paragraph()
    p.add_run('Duration: 4 Weeks').bold = True

    phase1_tasks = [
        'Project setup with Next.js 16, TypeScript, and Tailwind CSS',
        'Database schema design and Prisma setup with PostgreSQL',
        'Authentication system with JWT and session management',
        'User management with RBAC (roles and permissions)',
        'Basic UI component library setup with Radix UI',
        'Docker configuration for development environment',
        'CI/CD pipeline setup with Vercel',
    ]

    p = doc.add_paragraph()
    p.add_run('Deliverables:').bold = True
    for task in phase1_tasks:
        doc.add_paragraph(task, style='List Bullet')

    # Phase 2
    doc.add_heading('Phase 2: Parking Management Core', level=2)
    p = doc.add_paragraph()
    p.add_run('Duration: 6 Weeks').bold = True

    phase2_tasks = [
        'Parking lot CRUD with multi-venue support',
        'Zone and floor management',
        'Slot management with bulk creation and visual mapping',
        'Token generation system (QR, RFID, ANPR)',
        'Smart slot allocation algorithm',
        'Entry/exit workflow implementation',
        'Gate control integration (API-based)',
        'Display board integration',
    ]

    p = doc.add_paragraph()
    p.add_run('Deliverables:').bold = True
    for task in phase2_tasks:
        doc.add_paragraph(task, style='List Bullet')

    # Phase 3
    doc.add_heading('Phase 3: AI Pipeline Integration', level=2)
    p = doc.add_paragraph()
    p.add_run('Duration: 5 Weeks').bold = True

    phase3_tasks = [
        'OpenVINO environment setup and configuration',
        'YOLOv8 model integration for vehicle detection',
        'License plate detection and recognition models',
        'Camera management with RTSP/ONVIF support',
        'Real-time video stream processing',
        'Detection event pipeline with MQTT',
        'Occupancy detection with confidence scoring',
        'AI-to-backend communication via secure API',
    ]

    p = doc.add_paragraph()
    p.add_run('Deliverables:').bold = True
    for task in phase3_tasks:
        doc.add_paragraph(task, style='List Bullet')

    # Phase 4
    doc.add_heading('Phase 4: Payment & Wallet System', level=2)
    p = doc.add_paragraph()
    p.add_run('Duration: 5 Weeks').bold = True

    phase4_tasks = [
        'Wallet system architecture and database schema',
        'Wallet creation with KYC levels',
        'Deposit and withdrawal functionality',
        'P2P transfer implementation',
        'Razorpay integration (UPI, Cards)',
        'Stripe integration (Cards, Digital wallets)',
        'Parking payment workflow',
        'Bank account linking with verification',
        'Transaction history and statements',
        'Refund processing',
    ]

    p = doc.add_paragraph()
    p.add_run('Deliverables:').bold = True
    for task in phase4_tasks:
        doc.add_paragraph(task, style='List Bullet')

    # Phase 5
    doc.add_heading('Phase 5: Real-time & Notifications', level=2)
    p = doc.add_paragraph()
    p.add_run('Duration: 3 Weeks').bold = True

    phase5_tasks = [
        'Socket.IO server setup for real-time updates',
        'Live dashboard with occupancy updates',
        'Email notification service (SMTP/Resend/SendGrid)',
        'SMS integration (Twilio/MSG91)',
        'In-app notification system',
        'Alert rules engine with configurable thresholds',
        'Webhook support for external integrations',
    ]

    p = doc.add_paragraph()
    p.add_run('Deliverables:').bold = True
    for task in phase5_tasks:
        doc.add_paragraph(task, style='List Bullet')

    # Phase 6
    doc.add_heading('Phase 6: Analytics & Reporting', level=2)
    p = doc.add_paragraph()
    p.add_run('Duration: 3 Weeks').bold = True

    phase6_tasks = [
        'Analytics data aggregation jobs',
        'Dashboard metrics and KPIs',
        'Historical trend analysis',
        'Revenue reports with breakdowns',
        'Occupancy analytics with heatmaps',
        'CSV and PDF export functionality',
        'Predictive analytics foundation',
    ]

    p = doc.add_paragraph()
    p.add_run('Deliverables:').bold = True
    for task in phase6_tasks:
        doc.add_paragraph(task, style='List Bullet')

    # Phase 7
    doc.add_heading('Phase 7: Public Interfaces & Testing', level=2)
    p = doc.add_paragraph()
    p.add_run('Duration: 3 Weeks').bold = True

    phase7_tasks = [
        'Public kiosk interface for self-service',
        'Find My Car portal for vehicle location',
        'API documentation portal',
        'Comprehensive unit and integration testing',
        'Load testing and performance optimization',
        'Security audit and penetration testing',
        'Bug fixes and refinements',
    ]

    p = doc.add_paragraph()
    p.add_run('Deliverables:').bold = True
    for task in phase7_tasks:
        doc.add_paragraph(task, style='List Bullet')

    # Phase 8
    doc.add_heading('Phase 8: Deployment & Launch', level=2)
    p = doc.add_paragraph()
    p.add_run('Duration: 2 Weeks').bold = True

    phase8_tasks = [
        'Production environment setup',
        'Database migration and seeding',
        'SSL/TLS certificate configuration',
        'Monitoring and alerting setup',
        'Documentation finalization',
        'User training materials',
        'Go-live support and stabilization',
    ]

    p = doc.add_paragraph()
    p.add_run('Deliverables:').bold = True
    for task in phase8_tasks:
        doc.add_paragraph(task, style='List Bullet')

    doc.add_page_break()

    # ========== 7. TEAM SETUP ==========
    doc.add_heading('7. Required Team Setup', level=1)

    doc.add_heading('7.1 Core Development Team', level=2)

    team_table = doc.add_table(rows=1, cols=4)
    team_table.style = 'Table Grid'

    header_cells = team_table.rows[0].cells
    team_headers = ['Role', 'Count', 'Experience', 'Key Responsibilities']
    for i, header in enumerate(team_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '1a365d')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    team_data = [
        ('Project Manager', '1', '5+ years', 'Project planning, stakeholder management, delivery oversight'),
        ('Tech Lead / Architect', '1', '7+ years', 'Architecture design, technical decisions, code reviews'),
        ('Senior Full-Stack Developer', '2', '4+ years', 'Frontend/backend development, API design, database'),
        ('Full-Stack Developer', '2', '2+ years', 'Feature development, bug fixes, testing'),
        ('AI/ML Engineer', '1', '3+ years', 'OpenVINO integration, model optimization, detection pipeline'),
        ('DevOps Engineer', '1', '3+ years', 'Infrastructure, CI/CD, Docker, monitoring'),
        ('UI/UX Designer', '1', '3+ years', 'Interface design, user experience, prototyping'),
        ('QA Engineer', '1', '3+ years', 'Test planning, automation, quality assurance'),
    ]

    for role, count, exp, resp in team_data:
        row = team_table.add_row()
        row.cells[0].text = role
        row.cells[1].text = count
        row.cells[2].text = exp
        row.cells[3].text = resp

    doc.add_heading('7.2 Required Skills Matrix', level=2)

    skills_table = doc.add_table(rows=1, cols=3)
    skills_table.style = 'Table Grid'

    header_cells = skills_table.rows[0].cells
    skill_headers = ['Area', 'Required Skills', 'Nice to Have']
    for i, header in enumerate(skill_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '2e7d32')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    skills_data = [
        ('Frontend', 'React, Next.js, TypeScript, Tailwind CSS', 'React Query, Zustand, Radix UI'),
        ('Backend', 'Node.js, Prisma, PostgreSQL, REST APIs', 'GraphQL, Redis, Socket.IO'),
        ('AI/ML', 'Python, OpenCV, Deep Learning', 'OpenVINO, YOLO, MQTT'),
        ('DevOps', 'Docker, Linux, CI/CD', 'Kubernetes, AWS, Nginx'),
        ('Database', 'PostgreSQL, SQL optimization', 'Redis, Database design'),
        ('Payment', 'Payment gateway integration', 'Razorpay, Stripe, PCI compliance'),
    ]

    for area, required, nice in skills_data:
        row = skills_table.add_row()
        row.cells[0].text = area
        row.cells[1].text = required
        row.cells[2].text = nice

    doc.add_heading('7.3 Team Structure', level=2)

    p = doc.add_paragraph()
    p.add_run('Total Team Size: 10 members').bold = True

    structure_items = [
        'Sprint-based Agile methodology with 2-week sprints',
        'Daily stand-ups and weekly sprint reviews',
        'Dedicated QA involvement from Phase 2 onwards',
        'AI/ML engineer engaged primarily in Phases 3-4',
        'DevOps support throughout with focus on Phases 1 and 8',
    ]

    for item in structure_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ========== 8. TIMELINE ==========
    doc.add_heading('8. Project Timeline', level=1)

    doc.add_heading('8.1 Overall Timeline Summary', level=2)

    p = doc.add_paragraph()
    p.add_run('Total Project Duration: 31 Weeks (~8 Months)').bold = True
    p.paragraph_format.space_after = Pt(12)

    timeline_table = doc.add_table(rows=1, cols=4)
    timeline_table.style = 'Table Grid'

    header_cells = timeline_table.rows[0].cells
    timeline_headers = ['Phase', 'Description', 'Duration', 'Cumulative']
    for i, header in enumerate(timeline_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '1a365d')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    timeline_data = [
        ('Phase 1', 'Foundation & Core Infrastructure', '4 weeks', 'Week 1-4'),
        ('Phase 2', 'Parking Management Core', '6 weeks', 'Week 5-10'),
        ('Phase 3', 'AI Pipeline Integration', '5 weeks', 'Week 11-15'),
        ('Phase 4', 'Payment & Wallet System', '5 weeks', 'Week 16-20'),
        ('Phase 5', 'Real-time & Notifications', '3 weeks', 'Week 21-23'),
        ('Phase 6', 'Analytics & Reporting', '3 weeks', 'Week 24-26'),
        ('Phase 7', 'Public Interfaces & Testing', '3 weeks', 'Week 27-29'),
        ('Phase 8', 'Deployment & Launch', '2 weeks', 'Week 30-31'),
    ]

    for phase, desc, duration, cumulative in timeline_data:
        row = timeline_table.add_row()
        row.cells[0].text = phase
        row.cells[1].text = desc
        row.cells[2].text = duration
        row.cells[3].text = cumulative

    doc.add_heading('8.2 Milestone Schedule', level=2)

    milestones = [
        ('M1: Foundation Complete', 'Week 4', 'Authentication, user management, basic UI'),
        ('M2: Core Parking Ready', 'Week 10', 'Full parking management workflow'),
        ('M3: AI Integration Complete', 'Week 15', 'Vehicle and plate detection operational'),
        ('M4: Payment System Live', 'Week 20', 'All payment methods functional'),
        ('M5: Real-time Features', 'Week 23', 'Live updates and notifications'),
        ('M6: Analytics Dashboard', 'Week 26', 'Full reporting capabilities'),
        ('M7: Beta Release', 'Week 29', 'Feature-complete for testing'),
        ('M8: Production Launch', 'Week 31', 'Go-live with support'),
    ]

    milestone_table = doc.add_table(rows=1, cols=3)
    milestone_table.style = 'Table Grid'

    header_cells = milestone_table.rows[0].cells
    milestone_headers = ['Milestone', 'Target', 'Deliverables']
    for i, header in enumerate(milestone_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '2e7d32')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for name, target, deliverables in milestones:
        row = milestone_table.add_row()
        row.cells[0].text = name
        row.cells[1].text = target
        row.cells[2].text = deliverables

    doc.add_page_break()

    # ========== 9. API DOCUMENTATION ==========
    doc.add_heading('9. API Documentation', level=1)

    doc.add_heading('9.1 API Overview', level=2)

    p = doc.add_paragraph()
    p.add_run('SPARKING provides a comprehensive RESTful API with 60+ endpoints organized by domain. All endpoints follow consistent patterns and return JSON responses.')
    p.paragraph_format.space_after = Pt(12)

    api_overview = [
        ('Base URL', 'https://api.sparking.com/api or /api (relative)'),
        ('Authentication', 'JWT Bearer token in Authorization header or HttpOnly cookie'),
        ('Content-Type', 'application/json'),
        ('Error Format', '{ "error": "message", "code": "ERROR_CODE" }'),
    ]

    for key, value in api_overview:
        p = doc.add_paragraph()
        p.add_run(f'{key}: ').bold = True
        p.add_run(value)

    # Authentication APIs
    doc.add_heading('9.2 Authentication APIs', level=2)

    auth_table = doc.add_table(rows=1, cols=4)
    auth_table.style = 'Table Grid'

    header_cells = auth_table.rows[0].cells
    api_headers = ['Method', 'Endpoint', 'Description', 'Auth Required']
    for i, header in enumerate(api_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '1565c0')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    auth_apis = [
        ('POST', '/api/auth/login', 'User login with email/password', 'No'),
        ('POST', '/api/auth/logout', 'Terminate current session', 'Yes'),
        ('GET', '/api/auth/me', 'Get current user details', 'Yes'),
    ]

    for method, endpoint, desc, auth in auth_apis:
        row = auth_table.add_row()
        row.cells[0].text = method
        row.cells[1].text = endpoint
        row.cells[2].text = desc
        row.cells[3].text = auth

    # Parking Lot APIs
    doc.add_heading('9.3 Parking Lot APIs', level=2)

    lot_table = doc.add_table(rows=1, cols=4)
    lot_table.style = 'Table Grid'

    header_cells = lot_table.rows[0].cells
    for i, header in enumerate(api_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '2e7d32')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    lot_apis = [
        ('GET', '/api/parking-lots', 'List all parking lots with stats', 'Yes'),
        ('POST', '/api/parking-lots', 'Create new parking lot', 'Admin'),
        ('GET', '/api/parking-lots/{id}', 'Get parking lot details', 'Yes'),
        ('PATCH', '/api/parking-lots/{id}', 'Update parking lot', 'Admin'),
        ('GET', '/api/parking-lots/{id}/stats', 'Real-time statistics', 'Yes'),
    ]

    for method, endpoint, desc, auth in lot_apis:
        row = lot_table.add_row()
        row.cells[0].text = method
        row.cells[1].text = endpoint
        row.cells[2].text = desc
        row.cells[3].text = auth

    # Token APIs
    doc.add_heading('9.4 Token Management APIs', level=2)

    token_table = doc.add_table(rows=1, cols=4)
    token_table.style = 'Table Grid'

    header_cells = token_table.rows[0].cells
    for i, header in enumerate(api_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], 'f57c00')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    token_apis = [
        ('GET', '/api/tokens', 'List tokens with filters', 'Yes'),
        ('POST', '/api/tokens', 'Create entry token', 'Yes'),
        ('GET', '/api/tokens/{id}', 'Get token details', 'Yes'),
        ('PATCH', '/api/tokens/{id}', 'Update token (mark exit)', 'Yes'),
        ('DELETE', '/api/tokens/{id}', 'Cancel token', 'Admin'),
    ]

    for method, endpoint, desc, auth in token_apis:
        row = token_table.add_row()
        row.cells[0].text = method
        row.cells[1].text = endpoint
        row.cells[2].text = desc
        row.cells[3].text = auth

    # Payment APIs
    doc.add_heading('9.5 Payment & Wallet APIs', level=2)

    payment_table = doc.add_table(rows=1, cols=4)
    payment_table.style = 'Table Grid'

    header_cells = payment_table.rows[0].cells
    for i, header in enumerate(api_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '7b1fa2')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    payment_apis = [
        ('GET', '/api/wallet', 'Get user wallets', 'Yes'),
        ('POST', '/api/wallet', 'Create new wallet', 'Yes'),
        ('GET', '/api/wallet/{id}/balance', 'Get wallet balance', 'Yes'),
        ('POST', '/api/payments/parking', 'Pay for parking', 'Yes'),
        ('POST', '/api/payments/deposit', 'Deposit to wallet', 'Yes'),
        ('POST', '/api/payments/withdraw', 'Withdraw to bank', 'Yes'),
        ('POST', '/api/payments/transfer', 'P2P transfer', 'Yes'),
        ('GET', '/api/transactions', 'Transaction history', 'Yes'),
    ]

    for method, endpoint, desc, auth in payment_apis:
        row = payment_table.add_row()
        row.cells[0].text = method
        row.cells[1].text = endpoint
        row.cells[2].text = desc
        row.cells[3].text = auth

    # Analytics APIs
    doc.add_heading('9.6 Analytics APIs', level=2)

    analytics_table = doc.add_table(rows=1, cols=4)
    analytics_table.style = 'Table Grid'

    header_cells = analytics_table.rows[0].cells
    for i, header in enumerate(api_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '00796b')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    analytics_apis = [
        ('GET', '/api/analytics', 'Overview metrics', 'Yes'),
        ('GET', '/api/analytics?type=occupancy', 'Occupancy trends', 'Yes'),
        ('GET', '/api/analytics?type=revenue', 'Revenue analysis', 'Yes'),
        ('GET', '/api/analytics?type=traffic', 'Entry/exit patterns', 'Yes'),
        ('GET', '/api/analytics/predictive', 'Forecasting', 'Yes'),
        ('GET', '/api/reports/export', 'Export CSV/PDF', 'Yes'),
    ]

    for method, endpoint, desc, auth in analytics_apis:
        row = analytics_table.add_row()
        row.cells[0].text = method
        row.cells[1].text = endpoint
        row.cells[2].text = desc
        row.cells[3].text = auth

    # Real-time APIs
    doc.add_heading('9.7 Real-time & Detection APIs', level=2)

    realtime_table = doc.add_table(rows=1, cols=4)
    realtime_table.style = 'Table Grid'

    header_cells = realtime_table.rows[0].cells
    for i, header in enumerate(api_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], 'c62828')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    realtime_apis = [
        ('POST', '/api/realtime/detection', 'AI detection event (from pipeline)', 'API Key'),
        ('GET', '/api/cameras', 'List cameras', 'Yes'),
        ('GET', '/api/cameras/{id}/stream', 'RTSP stream proxy', 'Yes'),
        ('GET', '/api/metrics', 'System metrics', 'Yes'),
        ('WS', '/socket.io', 'Real-time updates', 'Yes'),
    ]

    for method, endpoint, desc, auth in realtime_apis:
        row = realtime_table.add_row()
        row.cells[0].text = method
        row.cells[1].text = endpoint
        row.cells[2].text = desc
        row.cells[3].text = auth

    doc.add_page_break()

    # ========== 10. DATABASE SCHEMA ==========
    doc.add_heading('10. Database Schema', level=1)

    doc.add_heading('10.1 Schema Overview', level=2)

    p = doc.add_paragraph()
    p.add_run('The SPARKING database consists of 27 interconnected models managed by Prisma ORM. The schema is designed for scalability, data integrity, and efficient querying.')
    p.paragraph_format.space_after = Pt(12)

    schema_stats = [
        ('Total Models', '27'),
        ('Total Fields', '250+'),
        ('Primary Database', 'PostgreSQL 16'),
        ('ORM', 'Prisma 7.2.0'),
        ('ID Strategy', 'CUID (collision-resistant unique identifiers)'),
    ]

    for key, value in schema_stats:
        p = doc.add_paragraph()
        p.add_run(f'{key}: ').bold = True
        p.add_run(value)

    # Core Models
    doc.add_heading('10.2 Core Models', level=2)

    # Organization
    doc.add_heading('Organization', level=3)
    org_table = doc.add_table(rows=1, cols=3)
    org_table.style = 'Table Grid'

    header_cells = org_table.rows[0].cells
    schema_headers = ['Field', 'Type', 'Description']
    for i, header in enumerate(schema_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '37474f')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    org_fields = [
        ('id', 'String @id @default(cuid())', 'Primary key'),
        ('name', 'String', 'Organization name'),
        ('slug', 'String @unique', 'URL-friendly identifier'),
        ('email', 'String?', 'Contact email'),
        ('phone', 'String?', 'Contact phone'),
        ('address', 'String?', 'Physical address'),
        ('logo', 'String?', 'Logo URL'),
        ('isActive', 'Boolean @default(true)', 'Active status'),
        ('createdAt', 'DateTime @default(now())', 'Creation timestamp'),
        ('updatedAt', 'DateTime @updatedAt', 'Last update'),
    ]

    for field, type_, desc in org_fields:
        row = org_table.add_row()
        row.cells[0].text = field
        row.cells[1].text = type_
        row.cells[2].text = desc

    # User
    doc.add_heading('User', level=3)
    user_table = doc.add_table(rows=1, cols=3)
    user_table.style = 'Table Grid'

    header_cells = user_table.rows[0].cells
    for i, header in enumerate(schema_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '37474f')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    user_fields = [
        ('id', 'String @id @default(cuid())', 'Primary key'),
        ('email', 'String @unique', 'Login email'),
        ('password', 'String', 'Hashed password (bcrypt)'),
        ('name', 'String', 'Display name'),
        ('role', 'UserRole @default(VIEWER)', 'SUPER_ADMIN, ADMIN, OPERATOR, AUDITOR, VIEWER'),
        ('status', 'UserStatus @default(ACTIVE)', 'ACTIVE, INACTIVE, SUSPENDED'),
        ('organizationId', 'String?', 'Foreign key to Organization'),
        ('maxSessions', 'Int @default(5)', 'Concurrent session limit'),
        ('createdAt', 'DateTime @default(now())', 'Creation timestamp'),
    ]

    for field, type_, desc in user_fields:
        row = user_table.add_row()
        row.cells[0].text = field
        row.cells[1].text = type_
        row.cells[2].text = desc

    # ParkingLot
    doc.add_heading('ParkingLot', level=3)
    lot_schema_table = doc.add_table(rows=1, cols=3)
    lot_schema_table.style = 'Table Grid'

    header_cells = lot_schema_table.rows[0].cells
    for i, header in enumerate(schema_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '37474f')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    lot_fields = [
        ('id', 'String @id @default(cuid())', 'Primary key'),
        ('name', 'String', 'Parking lot name'),
        ('type', 'VenueType', 'AIRPORT, MALL, HOSPITAL, STADIUM, etc.'),
        ('address', 'String', 'Physical address'),
        ('city', 'String', 'City'),
        ('latitude', 'Float?', 'GPS latitude'),
        ('longitude', 'Float?', 'GPS longitude'),
        ('totalCapacity', 'Int', 'Total slot count'),
        ('operatingHours', 'Json?', 'Operating schedule'),
        ('contactPhone', 'String?', 'Contact number'),
        ('isActive', 'Boolean @default(true)', 'Active status'),
        ('organizationId', 'String', 'Foreign key to Organization'),
    ]

    for field, type_, desc in lot_fields:
        row = lot_schema_table.add_row()
        row.cells[0].text = field
        row.cells[1].text = type_
        row.cells[2].text = desc

    # Slot
    doc.add_heading('Slot', level=3)
    slot_table = doc.add_table(rows=1, cols=3)
    slot_table.style = 'Table Grid'

    header_cells = slot_table.rows[0].cells
    for i, header in enumerate(schema_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '37474f')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    slot_fields = [
        ('id', 'String @id @default(cuid())', 'Primary key'),
        ('slotNumber', 'String', 'Display number (e.g., A-001)'),
        ('type', 'SlotType', 'STANDARD, COMPACT, LARGE, EV_CHARGING, etc.'),
        ('status', 'SlotStatus @default(AVAILABLE)', 'AVAILABLE, OCCUPIED, RESERVED, MAINTENANCE'),
        ('isOccupied', 'Boolean @default(false)', 'Quick occupancy check'),
        ('vehicleType', 'VehicleType?', 'CAR, SUV, MOTORCYCLE, etc.'),
        ('positionX', 'Float?', 'Visual map X coordinate'),
        ('positionY', 'Float?', 'Visual map Y coordinate'),
        ('rotation', 'Float? @default(0)', 'Visual rotation angle'),
        ('hasEvCharger', 'Boolean @default(false)', 'EV charging availability'),
        ('hasRoof', 'Boolean @default(false)', 'Covered parking'),
        ('isAccessible', 'Boolean @default(false)', 'Handicap accessible'),
        ('zoneId', 'String', 'Foreign key to Zone'),
        ('parkingLotId', 'String', 'Foreign key to ParkingLot'),
    ]

    for field, type_, desc in slot_fields:
        row = slot_table.add_row()
        row.cells[0].text = field
        row.cells[1].text = type_
        row.cells[2].text = desc

    # Token
    doc.add_heading('Token', level=3)
    token_schema_table = doc.add_table(rows=1, cols=3)
    token_schema_table.style = 'Table Grid'

    header_cells = token_schema_table.rows[0].cells
    for i, header in enumerate(schema_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '37474f')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    token_fields = [
        ('id', 'String @id @default(cuid())', 'Primary key'),
        ('tokenNumber', 'String @unique', 'Unique token identifier'),
        ('type', 'TokenType', 'QR_CODE, RFID, BARCODE, ANPR, MANUAL'),
        ('status', 'TokenStatus @default(ACTIVE)', 'ACTIVE, COMPLETED, CANCELLED, EXPIRED'),
        ('qrCode', 'String?', 'Generated QR code data'),
        ('licensePlate', 'String?', 'Vehicle license plate'),
        ('entryTime', 'DateTime @default(now())', 'Entry timestamp'),
        ('exitTime', 'DateTime?', 'Exit timestamp'),
        ('duration', 'Int?', 'Parking duration in minutes'),
        ('amount', 'BigInt?', 'Charged amount in paise'),
        ('isPaid', 'Boolean @default(false)', 'Payment status'),
        ('slotId', 'String?', 'Foreign key to Slot'),
        ('vehicleId', 'String?', 'Foreign key to Vehicle'),
        ('parkingLotId', 'String', 'Foreign key to ParkingLot'),
    ]

    for field, type_, desc in token_fields:
        row = token_schema_table.add_row()
        row.cells[0].text = field
        row.cells[1].text = type_
        row.cells[2].text = desc

    # Wallet
    doc.add_heading('Wallet', level=3)
    wallet_table = doc.add_table(rows=1, cols=3)
    wallet_table.style = 'Table Grid'

    header_cells = wallet_table.rows[0].cells
    for i, header in enumerate(schema_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '37474f')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    wallet_fields = [
        ('id', 'String @id @default(cuid())', 'Primary key'),
        ('type', 'WalletType', 'PERSONAL, BUSINESS, MERCHANT'),
        ('balance', 'BigInt @default(0)', 'Balance in paise (smallest unit)'),
        ('currency', 'String @default("INR")', 'Currency code'),
        ('isActive', 'Boolean @default(true)', 'Active status'),
        ('kycLevel', 'KycLevel @default(NONE)', 'NONE, BASIC, INTERMEDIATE, FULL'),
        ('dailyLimit', 'BigInt?', 'Daily transaction limit'),
        ('monthlyLimit', 'BigInt?', 'Monthly transaction limit'),
        ('perTransactionLimit', 'BigInt?', 'Single transaction limit'),
        ('userId', 'String?', 'Foreign key to User'),
        ('parkingLotId', 'String?', 'Foreign key (for merchant wallets)'),
    ]

    for field, type_, desc in wallet_fields:
        row = wallet_table.add_row()
        row.cells[0].text = field
        row.cells[1].text = type_
        row.cells[2].text = desc

    doc.add_heading('10.3 Model Relationships', level=2)

    relationships = [
        'Organization → Users, ParkingLots, Settings (one-to-many)',
        'ParkingLot → Zones, Slots, Tokens, Cameras, Gates, Displays (one-to-many)',
        'Zone → Slots, PricingRules (one-to-many)',
        'Slot → SlotOccupancy, Tokens (one-to-many)',
        'User → Sessions, Wallets, AuditLogs, ParkingLotAssignments (one-to-many)',
        'Token → Transaction, Vehicle, Slot (many-to-one)',
        'Wallet → WalletTransactions, BankAccounts, Payments (one-to-many)',
        'Camera → DetectionEvents (one-to-many)',
    ]

    for rel in relationships:
        doc.add_paragraph(rel, style='List Bullet')

    doc.add_heading('10.4 Indexes & Optimizations', level=2)

    indexes = [
        ('Slot', 'parkingLotId, zoneId, status, isOccupied'),
        ('Token', 'parkingLotId, status, entryTime, licensePlate'),
        ('WalletTransaction', 'walletId, type, status, createdAt'),
        ('DetectionEvent', 'cameraId, type, timestamp'),
        ('AuditLog', 'userId, entityType, createdAt'),
    ]

    for model, idx in indexes:
        p = doc.add_paragraph()
        p.add_run(f'{model}: ').bold = True
        p.add_run(idx)

    # ========== FOOTER ==========
    doc.add_page_break()

    doc.add_heading('Document Information', level=1)

    footer_info = [
        ('Document Version', '1.0'),
        ('Created Date', 'January 2026'),
        ('Last Updated', 'January 2026'),
        ('Prepared By', 'Development Team'),
        ('Confidentiality', 'This document contains proprietary information.'),
    ]

    for key, value in footer_info:
        p = doc.add_paragraph()
        p.add_run(f'{key}: ').bold = True
        p.add_run(value)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.add_run('For any questions or clarifications regarding this Statement of Work, please contact the project team.')

    # Save document
    doc_path = os.path.join(OUTPUT_DIR, 'SPARKING_Statement_of_Work.docx')
    doc.save(doc_path)
    print(f"Created: {doc_path}")
    return doc_path


def main():
    print("Generating SPARKING Statement of Work...")
    print("=" * 50)

    # Create diagrams
    print("\n[1/3] Creating Data Flow Diagram...")
    data_flow_path = create_data_flow_diagram()

    print("\n[2/3] Creating Software Architecture Diagram...")
    architecture_path = create_architecture_diagram()

    print("\n[3/3] Creating Statement of Work Document...")
    doc_path = create_sow_document(data_flow_path, architecture_path)

    print("\n" + "=" * 50)
    print("Generation Complete!")
    print(f"\nOutput Directory: {OUTPUT_DIR}")
    print("\nGenerated Files:")
    print(f"  1. {os.path.basename(data_flow_path)}")
    print(f"  2. {os.path.basename(architecture_path)}")
    print(f"  3. {os.path.basename(doc_path)}")


if __name__ == "__main__":
    main()
